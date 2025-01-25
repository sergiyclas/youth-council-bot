from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
import os
from datetime import datetime

from bot.common.infinitiveConverter import convert_to_infinitive, capitalize_first_word
from bot.common.mini_libs import months_uk, questions

def set_page_margins(document):
    """
    Налаштовує поля документа згідно з вимогами.
    Верхнє: 1,59 см
    Нижнє: 1,88 см
    Ліве: 2,5 см
    Праве: 1,5 см
    """
    section = document.sections[0]
    section.page_height = Cm(29.7)  # Висота для аркуша А4
    section.page_width = Cm(21.0)  # Ширина для аркуша А4
    section.orientation = WD_ORIENTATION.PORTRAIT

    # Поля
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)


async def generate_protocol(session_code, db):
    """
    Генерує протокол для завершеної сесії.
    """
    # Отримуємо дані про сесію
    session = await db.get_session_by_code(session_code)
    if not session:
        raise ValueError("Сесія не знайдена.")

    admin_id = session.admin_id
    participants = await db.get_session_participants_with_names(session_code) or []
    agenda = await db.get_session_agenda(session_code) or []
    voting_results = await db.get_all_vote_results(session_code) or {}

    print(voting_results)

    youth_council_info = await db.get_full_youth_council_info(admin_id)
    protocol_info = await db.get_session_details(session_code)

    council_name = youth_council_info.get("name", "______________________________________________________")
    if council_name.lower().strip()[:14] == 'молодіжна рада':
        council_name = "Молодіжної ради " + council_name[15:]

    city = youth_council_info.get("city", "______________")
    head = youth_council_info.get("head", "___________________________")
    secretary = youth_council_info.get("secretary", "___________________________")
    number = protocol_info.get("number", "____")
    session_type = protocol_info.get('session_type', "______________")

    try:
        date = session.date.strftime("%Y_%m_%d_%H_%M")
    except Exception as e:
        date = session.date

    file_name = f"{date}_Протокол_{number}.docx"
    file_path = os.path.join("protocols", file_name)

    document = Document()
    set_page_margins(document)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run(f"Протокол № {number}\n{session_type} засідання {council_name}")
    title_run.bold = True
    title_run.font.size = Pt(14)

    date_obj = datetime.strptime(date, "%Y_%m_%d_%H_%M")
    date_for_title = f"{date_obj.day} {months_uk[date_obj.month]} {date_obj.year}"

    font_size = 14
    max_width = 113 - len(city)

    text_with_spaces = f"{city}{' ' * (max_width - len(date_for_title) - 4)}{date_for_title} р."

    paragraph = document.add_paragraph()
    run = paragraph.add_run(text_with_spaces)
    run.font.size = Pt(font_size)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    present_paragraph = document.add_paragraph()
    present_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    present_paragraph.add_run(f"Присутні: {len(participants)} (додаток 1)").bold = True

    invited_paragraph = document.add_paragraph()
    invited_paragraph.add_run("Запрошені: ").bold = True
    invited_paragraph.add_run("-")

    agenda_paragraph = document.add_paragraph()
    agenda_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    agenda_paragraph.add_run("Порядок денний:").bold = True

    for i, item in enumerate(agenda, start=1):
        par = document.add_paragraph(style='List Number')
        run = par.add_run(f"{item}")
        run.bold = True
    par = document.add_paragraph(style='List Number')
    run = par.add_run(f"Різне")
    run.bold = True

    for i, (question, results) in enumerate(voting_results.items(), start=1):
        question_inf = convert_to_infinitive(question)
        paragraph = document.add_paragraph(style='Normal')
        run = paragraph.add_run(f"{i}. По {questions[i]} питанню порядку денного слухали")
        run.bold = True

        proposed_name = await db.get_proposed_name(session_code, question) or "_________________"
        proposed_rv = await db.get_name_rv(admin_id, proposed_name)
        proposer_text = proposed_rv.name_rv if proposed_rv and proposed_rv.name_rv else proposed_name

        document.add_paragraph(
            f"{proposer_text}, який запропонував {question_inf}", style='Normal'
        )

        not_vote = ''
        if results['not_voted'] >= 1:
            not_vote = f", \"НЕ ГОЛОСУВАЛИ\" - {results['not_voted']}"

        paragraph = document.add_paragraph(style='Normal')
        run = paragraph.add_run(
            f"Проголосували: \"ЗА\" - {results['for']}, \"ПРОТИ\" - {results['against']}, \"УТРИМАЛИСЬ\" - {results['abstain']}{not_vote}"
        )
        run.bold = True

        paragraph = document.add_paragraph(style='Normal')
        run = paragraph.add_run(f"Ухвалили:")
        run.bold = True

        paragraph = document.add_paragraph(style='Normal')
        run = paragraph.add_run(f"{capitalize_first_word(question_inf)}")

    paragraph = document.add_paragraph(style='Normal')
    run = paragraph.add_run(f"{len(voting_results.items()) + 1}. Питання різне - не піднімалося")
    run.bold = True

    document.add_paragraph("\n")

    paragraph = document.add_paragraph(style='Normal')
    run = paragraph.add_run("Головуючий засідання: ")
    run.bold = True
    run = paragraph.add_run(f"{head} _______________")

    document.add_paragraph("\n\n")

    paragraph = document.add_paragraph(style='Normal')
    run = paragraph.add_run("Секретар засідання: ")
    run.bold = True
    run = paragraph.add_run(f"{secretary} ______________")

    os.makedirs("protocols", exist_ok=True)
    document.save(file_path)

    return file_path


async def generate_attendance_list_full(session_code, db):
    """
    Генерує анкету присутності для сесії.

    :param session_code: Код сесії.
    :param db: Об'єкт бази даних.
    :return: Шлях до збереженого файлу анкети.
    """
    session = await db.get_session_by_code(session_code)
    if not session:
        raise ValueError("Сесія не знайдена.")

    admin_id = session.admin_id
    participants = await db.get_session_participants_with_names(session_code) or []

    youth_council_info = await db.get_full_youth_council_info(admin_id)
    protocol_info = await db.get_session_details(session_code)

    council_name = youth_council_info.get("name", "______________________________________________________")
    if council_name.lower().strip()[:14] == 'молодіжна рада':
        council_name = "Молодіжної ради " + council_name[19:]

    region = youth_council_info.get("region", "___________________________")
    number = protocol_info.get("number", "____")
    session_type = protocol_info.get('session_type', "______________")

    # Формуємо ім'я файлу додатка
    try:
        date = session.date.strftime("%Y_%m_%d_%H_%M")
    except Exception as e:
        date = session.date

    file_name = f"{date}_Додаток_присутності_{number}.docx"
    file_path = os.path.join("protocols", file_name)

    # Створюємо документ
    document = Document()
    set_page_margins(document)

    # Додаток 1
    appendix = document.add_paragraph()
    appendix.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    appendix_run = appendix.add_run("Додаток 1")
    appendix_run.font.name = "Times New Roman"
    appendix_run.font.size = Pt(14)
    appendix_run.bold = True

    # Заголовок
    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    date_obj = datetime.strptime(date, "%Y_%m_%d_%H_%M")
    date_for_title = f"{date_obj.day} {months_uk[date_obj.month]} {date_obj.year}"

    title_run = title.add_run(
        f"Реєстраційна анкета {session_type} засідання №{number}, {date_for_title} року\n"
        f"{council_name} {region}"
    )
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(14)
    title_run.bold = False

    # Таблиця присутності
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Заголовки таблиці
    header_cells = table.rows[0].cells
    header_cells[0].text = "№ п/п"
    header_cells[1].text = "Прізвище, ім'я"
    header_cells[2].text = "Присутність"

    # Форматування заголовків таблиці
    for header_cell in header_cells:
        for paragraph in header_cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)
                run.bold = True

    # Заповнення таблиці
    for index, participant in enumerate(participants, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(index)
        row_cells[1].text = participant["name"]
        row_cells[2].text = "+"

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(14)

    document.save(file_path)

    return file_path